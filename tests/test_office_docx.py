"""M4 DOCX-builder tests: the constrained DOCX builder fills styled blocks from a doc spec."""
import pytest

pytest.importorskip("docx")

from tools.office_docx import build_doc  # noqa: E402

_SPEC = {
    "title": "Operations Report",
    "subtitle": "Q2 2026",
    "theme": "light",
    "blocks": [
        {"type": "heading", "text": "Summary", "level": 1},
        {"type": "paragraph", "text": "Things went well this quarter."},
        {"type": "bullets", "items": ["Shipped X", "Hired Y", "Cut cost Z"]},
        {"type": "heading", "text": "Numbers", "level": 2},
        {"type": "table", "headers": ["Metric", "Value"], "rows": [["Revenue", "1.2M"], ["Churn", "2%"]]},
        {"type": "pagebreak"},
        {"type": "paragraph", "text": "Appendix."},
    ],
}


def test_build_all_block_types(tmp_path):
    out = tmp_path / "report.docx"
    res = build_doc(_SPEC, str(out))
    assert out.exists() and out.stat().st_size > 0
    assert res["blocks"] == 7
    assert res["warnings"] == []
    from docx import Document

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Operations Report" in text and "Summary" in text and "Appendix." in text
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Metric"


def test_unknown_block_falls_back_with_warning(tmp_path):
    res = build_doc({"blocks": [{"type": "nope", "text": "x"}]}, str(tmp_path / "x.docx"))
    assert any("unknown type" in w for w in res["warnings"])


def test_bad_spec_raises(tmp_path):
    with pytest.raises(ValueError):
        build_doc(["not", "a", "dict"], str(tmp_path / "x.docx"))
