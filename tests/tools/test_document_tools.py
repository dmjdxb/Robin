"""Tests for the "Ask your document" tool (tools/document_tools.py).

Fixtures are generated programmatically with the real parser libraries so there
are no binary blobs in git and the extraction round-trips end to end. Each
format test is guarded with ``importorskip`` — they run in CI (where the
``document`` extra installs) and skip cleanly where the libs are absent.

The format-independent helpers (anchors, budgeting, outline/section, rendering)
are tested WITHOUT any heavy libraries, so they always run.
"""

from __future__ import annotations

import pytest

import tools.document_tools as dt


# ─────────────────────────────────────────────────────────────────────────────
# Registration / wiring (no heavy libs)
# ─────────────────────────────────────────────────────────────────────────────
def test_ask_document_registered():
    from tools.registry import registry
    assert "ask_document" in registry.get_all_tool_names()
    entry = registry._tools["ask_document"]
    assert entry.toolset == "file"
    assert entry.max_result_size_chars == 100_000


def test_ask_document_in_core_and_document_toolset():
    import toolsets as ts
    assert "ask_document" in ts._HERMES_CORE_TOOLS
    assert "ask_document" in ts.TOOLSETS["file"]["tools"]
    assert ts.TOOLSETS["document"]["tools"] == ["ask_document", "read_file"]


def test_module_imports_without_heavy_libs():
    # The parser libs must be lazy — importing the module must not require them.
    assert dt.pypdf is None or dt.pypdf is not None  # tolerant: may be bound by a prior test
    # Schema is well-formed and carries the grounding contract.
    desc = dt.ASK_DOCUMENT_SCHEMA["description"].lower()
    assert "only from" in desc and "cite" in desc and "couldn't find" in desc.replace("’", "'")


# ─────────────────────────────────────────────────────────────────────────────
# Helpers (no heavy libs)
# ─────────────────────────────────────────────────────────────────────────────
def test_col_letter():
    assert dt._col_letter(1) == "A"
    assert dt._col_letter(26) == "Z"
    assert dt._col_letter(27) == "AA"
    assert dt._col_letter(28) == "AB"


def test_rows_to_md():
    md = dt._rows_to_md([["Name", "Qty"], ["Widget", "3"]])
    assert "| Name | Qty |" in md
    assert "|---|---|" in md
    assert "| Widget | 3 |" in md


def test_section_slice():
    chunks = [{"anchor": f"[p.{i}]", "text": "x"} for i in range(1, 6)]
    assert dt._slice_section(chunks, "2-4") == chunks[1:4]
    assert dt._slice_section(chunks, "1-1") == chunks[:1]
    assert dt._slice_section(chunks, "bad") is None
    # clamps out-of-range
    assert dt._slice_section(chunks, "3-99") == chunks[2:]


def test_render_full_framing():
    chunks = [{"anchor": "[p.1]", "text": "hello world"}]
    out = dt._render_full(chunks, {"type": "PDF", "unit": "pages", "count": 1}, "a.pdf")
    assert out.startswith("=== DOCUMENT: a.pdf ===")
    assert "[p.1]" in out and "hello world" in out
    assert out.rstrip().endswith("=== END DOCUMENT (a.pdf) ===")


def test_render_outline_has_section_ids():
    big = [{"anchor": f"[p.{i}]", "text": "x" * 30_000} for i in range(1, 6)]
    out = dt._render_outline(big, {"type": "PDF", "unit": "pages", "count": 5}, "big.pdf")
    assert "DOCUMENT OUTLINE" in out
    assert "id:" in out
    # the section ids must round-trip through _slice_section
    import re
    ids = re.findall(r"id:\s*([0-9]+-[0-9]+)", out)
    assert ids, out
    for sid in ids:
        assert dt._slice_section(big, sid) is not None


# ─────────────────────────────────────────────────────────────────────────────
# Error paths (no heavy libs)
# ─────────────────────────────────────────────────────────────────────────────
def test_missing_path():
    assert dt.ask_document_tool("").startswith("ERROR")


def test_not_found(tmp_path):
    assert "not found" in dt.ask_document_tool(str(tmp_path / "nope.pdf"))


def test_legacy_format_guidance(tmp_path):
    p = tmp_path / "old.doc"
    p.write_bytes(b"\x00legacy")
    out = dt.ask_document_tool(str(p))
    assert "re-save it as .docx" in out


def test_unsupported_format(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("hi")
    out = dt.ask_document_tool(str(p))
    assert "unsupported document type" in out and "read_file" in out


# ─────────────────────────────────────────────────────────────────────────────
# Per-format extraction (guarded — run when the libs are installed, e.g. in CI)
# ─────────────────────────────────────────────────────────────────────────────
def _docx_or_skip():
    docx = pytest.importorskip("docx")
    try:  # some environments ship a mismatched docx where docx.table is broken
        import docx.table  # noqa: F401
    except Exception as e:  # pragma: no cover - env-specific
        pytest.skip(f"installed python-docx is not usable: {e}")
    return docx


def test_docx_extraction(tmp_path):
    docx = _docx_or_skip()
    assert dt._ensure_docx()
    d = docx.Document()
    d.add_heading("Payment Terms", level=1)
    d.add_paragraph("Net-30 from invoice date.")
    p = tmp_path / "contract.docx"
    d.save(str(p))

    out = dt.ask_document_tool(str(p))
    assert "[§ Payment Terms]" in out
    assert "Net-30" in out
    assert "=== END DOCUMENT (contract.docx) ===" in out


def test_pptx_extraction(tmp_path):
    pptx = pytest.importorskip("pptx")
    assert dt._ensure_pptx()
    prs = pptx.Presentation()
    slide = prs.slides.add_slide(prs.slide_layouts[1])
    slide.shapes.title.text = "Revenue Forecast"
    slide.placeholders[1].text = "Q3 revenue up 12%."
    p = tmp_path / "deck.pptx"
    prs.save(str(p))

    out = dt.ask_document_tool(str(p))
    assert "[slide 1]" in out
    assert "Revenue Forecast" in out
    assert "Q3 revenue up 12%." in out


def test_xlsx_extraction_with_cell_anchors(tmp_path):
    openpyxl = pytest.importorskip("openpyxl")
    assert dt._ensure_xlsx()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws["A1"] = "Item"
    ws["B1"] = "Cost"
    ws["A2"] = "Rent"
    ws["B2"] = 5000
    p = tmp_path / "budget.xlsx"
    wb.save(str(p))

    out = dt.ask_document_tool(str(p))
    assert "[Sheet:Budget]" in out
    # column-letter + row-number framing makes cells citable as [Budget!B2]
    assert "| A | B |" in out
    assert "Rent" in out and "5000" in out


def test_pdf_extraction(tmp_path):
    # Build a tiny PDF with a text layer using reportlab if available; otherwise
    # skip (pypdf can read but we need a generator to create the fixture).
    pytest.importorskip("pypdf")
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas  # noqa: F401
    from reportlab.lib.pagesizes import letter

    p = tmp_path / "doc.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 720, "Termination requires thirty days notice.")
    c.showPage()
    c.drawString(72, 720, "Governing law is the DIFC.")
    c.showPage()
    c.save()

    out = dt.ask_document_tool(str(p))
    assert "[p.1]" in out and "[p.2]" in out
    assert "thirty days notice" in out
    assert "DIFC" in out


def test_oversized_pdf_returns_outline(tmp_path, monkeypatch):
    # Drive the OUTLINE path deterministically without a giant real PDF:
    # monkeypatch the extractor to return many large chunks.
    big = [{"anchor": f"[p.{i}]", "text": "x" * 20_000} for i in range(1, 11)]
    monkeypatch.setattr(dt, "_extract_pdf", lambda path: (big, {"type": "PDF", "unit": "pages", "count": 10}))
    monkeypatch.setattr(dt, "_ensure_pdf", lambda: True)
    p = tmp_path / "huge.pdf"
    p.write_bytes(b"%PDF-1.4 fake")

    out = dt.ask_document_tool(str(p))
    assert "DOCUMENT OUTLINE" in out

    # A section follow-up returns full text for that slice.
    import re
    sid = re.findall(r"id:\s*([0-9]+-[0-9]+)", out)[0]
    section_out = dt.ask_document_tool(str(p), section=sid)
    assert section_out.startswith("=== DOCUMENT: huge.pdf ===")
    assert "DOCUMENT OUTLINE" not in section_out
