"""Constrained DOCX builder — styled blocks, never freeform-blind.

The model supplies a document spec (title + ordered blocks); this builder owns the
styles, margins and spacing. Named styles (Title, Headings, Body) are set once per
theme so output is consistent. Pairs with the render_check vision gate (DOCX and PDF
both render). A report PDF is this DOCX converted via LibreOffice.

Doc spec (JSON):
  {
    "title": "...", "subtitle": "...", "theme": "light",
    "blocks": [
      {"type": "heading",   "text": "...", "level": 1},
      {"type": "paragraph", "text": "..."},
      {"type": "bullets",   "items": ["...", "..."]},
      {"type": "table",     "headers": ["A","B"], "rows": [["1","2"]]},
      {"type": "pagebreak"}
    ]
  }
"""
from __future__ import annotations

import json
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any

Document = Pt = Inches = RGBColor = None


def _ensure_docx() -> bool:
    global Document, Pt, Inches, RGBColor
    if Document is not None:
        return True

    def _imp() -> dict:
        from docx import Document as _D
        from docx.shared import Inches as _I
        from docx.shared import Pt as _Pt
        from docx.shared import RGBColor as _RGB

        return {"Document": _D, "Pt": _Pt, "Inches": _I, "RGBColor": _RGB}

    from tools.lazy_deps import ensure_and_bind

    return ensure_and_bind("document.docx", _imp, globals(), prompt=False)


@dataclass(frozen=True)
class Theme:
    ink: str
    body: str
    accent: str
    muted: str
    font: str = "Calibri"  # widely available; substituted metric-compatibly on Linux


THEMES: dict[str, Theme] = {
    "light": Theme(ink="1A1D21", body="2E3338", accent="2F8F6B", muted="6B7178"),
    "dark": Theme(ink="1A1D21", body="2E3338", accent="2F8F6B", muted="6B7178"),
}


def _rgb(h: str):
    return RGBColor.from_string(h)


def _style(doc, name: str, *, size: int, color: str, theme: Theme, bold: bool = False) -> None:
    try:
        st = doc.styles[name]
    except KeyError:
        return
    f = st.font
    f.name = theme.font
    f.size = Pt(size)
    f.bold = bold
    f.color.rgb = _rgb(color)


def _style_doc(doc, theme: Theme) -> None:
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)
    _style(doc, "Title", size=26, color=theme.ink, theme=theme, bold=True)
    _style(doc, "Heading 1", size=16, color=theme.accent, theme=theme, bold=True)
    _style(doc, "Heading 2", size=13, color=theme.ink, theme=theme, bold=True)
    _style(doc, "Normal", size=11, color=theme.body, theme=theme)


def _as_list(v: Any) -> list[str]:
    if v is None:
        return []
    if isinstance(v, str):
        return [v]
    return [str(x) for x in v]


# ── markdown → doc spec (deterministic; the fast path for "make a docx from this .md") ──
import re as _re

_EMPHASIS = [
    (_re.compile(r"!\[[^\]]*\]\([^)]*\)"), ""),          # drop image markdown (handled separately)
    (_re.compile(r"\[([^\]]+)\]\([^)]*\)"), r"\1"),       # [text](url) -> text
    (_re.compile(r"`([^`]*)`"), r"\1"),                    # `code` -> code
    (_re.compile(r"\*\*([^*]+)\*\*"), r"\1"),              # **bold** -> bold
    (_re.compile(r"\*([^*]+)\*"), r"\1"),                  # *italic* -> italic
    (_re.compile(r"__([^_]+)__"), r"\1"),
    (_re.compile(r"~~([^~]+)~~"), r"\1"),
]


def _strip_inline(s: str) -> str:
    out = s.strip()
    for pat, repl in _EMPHASIS:
        out = pat.sub(repl, out)
    return out.strip()


_IMG_RE = _re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)\s]+)(?:\s+\"[^\"]*\")?\)")


def markdown_to_spec(md: str, *, theme: str = "light", title: str | None = None) -> dict:
    """Parse markdown into a doc spec (title + ordered blocks) the builder consumes.

    Deterministic and dependency-free — no model reasoning needed to lay out a
    document from a markdown source. Supports: # headings, paragraphs, -/*/+ and
    numbered lists, GFM tables, ![alt](path) images, fenced code (as paragraphs),
    and blockquotes. Inline emphasis/links are flattened to clean text.
    """
    lines = md.replace("\r\n", "\n").split("\n")
    # strip YAML frontmatter
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                lines = lines[i + 1 :]
                break
    blocks: list[dict] = []
    para: list[str] = []
    bullets: list[str] = []
    doc_title = title

    def flush_para():
        nonlocal para
        if para:
            text = _strip_inline(" ".join(p.strip() for p in para).strip())
            if text:
                blocks.append({"type": "paragraph", "text": text})
            para = []

    def flush_bullets():
        nonlocal bullets
        if bullets:
            blocks.append({"type": "bullets", "items": bullets})
            bullets = []

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.strip()
        # fenced code block → collect verbatim as a paragraph
        if line.startswith("```"):
            flush_para(); flush_bullets()
            code: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            if code:
                blocks.append({"type": "paragraph", "text": "\n".join(code)})
            continue
        # standalone image
        m = _IMG_RE.match(line)
        if m:
            flush_para(); flush_bullets()
            blocks.append({"type": "image", "path": m.group("src"), "caption": _strip_inline(m.group("alt"))})
            i += 1
            continue
        # GFM table: a header row followed by a |---|---| divider
        if line.startswith("|") and i + 1 < n and _re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", lines[i + 1]) and "-" in lines[i + 1]:
            flush_para(); flush_bullets()
            def _cells(r: str) -> list[str]:
                return [_strip_inline(c) for c in r.strip().strip("|").split("|")]
            headers = _cells(line)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_cells(lines[i])); i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue
        # headings
        hm = _re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            flush_para(); flush_bullets()
            level = len(hm.group(1))
            text = _strip_inline(hm.group(2))
            if level == 1 and doc_title is None:
                doc_title = text
            else:
                blocks.append({"type": "heading", "text": text, "level": 1 if level <= 2 else 2})
            i += 1
            continue
        # list items
        lm = _re.match(r"^\s*(?:[-*+]|\d+[.)])\s+(.*)$", raw)
        if lm:
            flush_para()
            bullets.append(_strip_inline(lm.group(1)))
            i += 1
            continue
        # blockquote → paragraph
        if line.startswith(">"):
            flush_bullets()
            para.append(_strip_inline(line.lstrip(">").strip()))
            i += 1
            continue
        # blank line → paragraph/bullets break
        if not line:
            flush_para(); flush_bullets()
            i += 1
            continue
        # horizontal rule → ignore (avoid stray pagebreaks/blank pages)
        if _re.match(r"^([-*_])\1{2,}$", line):
            flush_para(); flush_bullets()
            i += 1
            continue
        # default: paragraph text
        flush_bullets()
        para.append(line)
        i += 1
    flush_para(); flush_bullets()
    spec: dict = {"theme": theme, "blocks": blocks}
    if doc_title:
        spec["title"] = doc_title
    return spec


def build_doc(spec: dict, out_path: str) -> dict:
    """Build a .docx from a doc spec. Returns {path, blocks, warnings}. Unknown block
    types become a plain paragraph with a warning — never crashes."""
    if not _ensure_docx():
        raise RuntimeError("python-docx unavailable — cannot build the document")
    if not isinstance(spec, dict):
        raise ValueError("doc spec must be an object")
    theme = THEMES.get(str(spec.get("theme", "light")), THEMES["light"])
    doc = Document()
    _style_doc(doc, theme)

    if spec.get("title"):
        doc.add_paragraph(str(spec["title"]), style="Title")
    if spec.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(str(spec["subtitle"]))
        run.font.size = Pt(13)
        run.font.color.rgb = _rgb(theme.muted)
        run.font.name = theme.font

    warnings: list[str] = []
    blocks = spec.get("blocks") or []
    for i, b in enumerate(blocks):
        if not isinstance(b, dict):
            warnings.append(f"block {i}: not an object, skipped")
            continue
        bt = str(b.get("type", "paragraph"))
        if bt == "heading":
            lvl = max(1, min(2, int(b.get("level", 1))))
            doc.add_paragraph(str(b.get("text", "")), style=f"Heading {lvl}")
        elif bt == "paragraph":
            doc.add_paragraph(str(b.get("text", "")), style="Normal")
        elif bt == "bullets":
            for item in _as_list(b.get("items")):
                doc.add_paragraph(item, style="List Bullet")
        elif bt == "table":
            headers = _as_list(b.get("headers"))
            rows = b.get("rows") or []
            ncols = len(headers) or (len(rows[0]) if rows else 0)
            if ncols:
                t = doc.add_table(rows=0, cols=ncols)
                try:
                    t.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001 — fall back if the style is unavailable
                    t.style = "Table Grid"
                if headers:
                    cells = t.add_row().cells
                    for c, h in zip(cells, headers):
                        c.text = str(h)
                for r in rows:
                    cells = t.add_row().cells
                    for c, val in zip(cells, _as_list(r)):
                        c.text = str(val)
        elif bt == "image":
            # Embed an illustration/picture file, centered, sized to the page
            # content width (the template owns sizing — the model never sets
            # raw EMU). An optional italic caption sits beneath it.
            img_path = str(b.get("path", "")).strip()
            if not img_path or not os.path.exists(img_path):
                warnings.append(f"block {i}: image not found: {img_path!r}")
            else:
                try:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    w = max(1.0, min(6.5, float(b.get("width_in", 6.0))))  # clamp to usable width
                    doc.add_picture(img_path, width=Inches(w))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = str(b.get("caption", "")).strip()
                    if cap:
                        p = doc.add_paragraph(cap)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.italic = True
                except Exception as e:  # noqa: BLE001 — never crash the whole doc on one image
                    warnings.append(f"block {i}: image embed failed: {e}")
        elif bt == "pagebreak":
            doc.add_page_break()
        else:
            warnings.append(f"block {i}: unknown type {bt!r} → paragraph")
            doc.add_paragraph(str(b.get("text", "")), style="Normal")

    doc.save(out_path)
    return {"path": out_path, "blocks": len(blocks), "warnings": warnings}


# ── in-process tool: build_document ────────────────────────────────────────────
# Registered as a TOOL (not a terminal script) so it runs in the Robin backend
# where python-docx auto-installs via lazy_deps. A terminal `python build_docx.py`
# runs the user's system Python, which has neither the `tools` package nor
# python-docx — the v1.2.17 install failure (ModuleNotFoundError: No module 'docx').
BUILD_DOCUMENT_SCHEMA = {
    "name": "build_document",
    "description": (
        "Build a polished .docx using designed templates that own all layout/styles/margins — "
        "never hand-write python-docx. FAST PATH: pass `markdown_path` (or `markdown` text) to "
        "convert an existing markdown source straight into a clean document in one call — no need "
        "to write a spec. Or pass a CONTENT `spec`: {title, subtitle?, theme:'light'|'dark', "
        "blocks:[...]}; block types: heading {text, level:1|2}, paragraph {text}, bullets {items}, "
        "table {headers, rows}, image {path, caption?, width_in?}, pagebreak. Embed illustrations "
        "with image blocks. Set verify=true to auto-run the render_check vision gate and get a "
        "pass/fail verdict back. Provide exactly one of markdown_path | markdown | spec."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "markdown_path": {"type": "string", "description": "Path to a .md/.txt source to convert (fast path)."},
            "markdown": {"type": "string", "description": "Markdown text to convert (fast path)."},
            "spec": {"type": "object", "description": "A content spec (see description) for full control."},
            "out_path": {"type": "string", "description": "Absolute output path ending in .docx"},
            "theme": {"type": "string", "description": "'light' (default) or 'dark' — used with markdown input."},
            "verify": {"type": "boolean", "description": "If true, auto-run render_check and include the verdict."},
        },
        "required": ["out_path"],
    },
}


def _spec_from_args(args: dict) -> dict | str:
    """Resolve a doc spec from markdown_path | markdown | spec. Returns spec dict or error str."""
    theme = str(args.get("theme") or "light")
    md_path = str(args.get("markdown_path") or "").strip()
    md_text = args.get("markdown")
    spec = args.get("spec")
    if md_path:
        try:
            md_text = Path(os.path.expanduser(md_path)).read_text(encoding="utf-8")
        except Exception as e:  # noqa: BLE001
            return f"could not read markdown_path: {e}"
    if isinstance(md_text, str) and md_text.strip():
        return markdown_to_spec(md_text, theme=theme)
    if isinstance(spec, str):
        try:
            spec = json.loads(spec)
        except Exception:
            return "spec must be a JSON object"
    if isinstance(spec, dict):
        spec.setdefault("theme", theme)
        return spec
    return "provide one of markdown_path, markdown, or spec"


async def _build_document_handler(args: dict, **_kw) -> str:
    out_path = str(args.get("out_path") or "").strip()
    if not out_path.lower().endswith(".docx"):
        return json.dumps({"error": "out_path must end in .docx"})
    spec = _spec_from_args(args)
    if isinstance(spec, str):
        return json.dumps({"error": spec})
    try:
        res = build_doc(spec, out_path)
    except Exception as e:  # noqa: BLE001
        return json.dumps({"error": f"build_document failed: {e}"})
    if args.get("verify"):
        try:
            from tools.document_qa import check_document
            res["render_check"] = await check_document(out_path)
        except Exception as e:  # noqa: BLE001
            res["render_check"] = {"error": f"verify failed: {e}"}
    return json.dumps(res)


# Top-level registration so discover_builtin_tools() (which only detects a top-level
# registry.register call) actually picks this up — a try/except-wrapped call is invisible
# to discovery, which is why the office tools were never registered/callable before.
from tools.registry import registry  # noqa: E402

registry.register(
    name="build_document",
    toolset="office",
    schema=BUILD_DOCUMENT_SCHEMA,
    handler=_build_document_handler,
    check_fn=lambda: True,  # python-docx lazy-installs on first use
    emoji="📄",
)
